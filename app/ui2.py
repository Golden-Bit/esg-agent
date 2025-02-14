import json
import os
import base64
import uuid
from typing import Any
import requests
import streamlit as st
import copy
import re
from docx import Document
from html2docx import html2docx

from scope3_form import render_scope3_form
from scope1_form import render_scope1_form
from scope2_form import render_scope2_form
from forms_ui import render_questionnaire, load_questions
from utilities import graph_notes_2, graph_notes_1_5, graph_notes_3_1, graph_notes_4_1, graph_notes_ESRS_index
from utilities import linee_guida_intro, linee_guida_1_1, linee_guida_1_2, linee_guida_1_3, linee_guida_2, linee_guida_3, linee_guida_1_4, linee_guida_1_5, linee_guida_2_1, linee_guida_2_2, linee_guida_2_3, linee_guida_2_4, linee_guida_2_5, linee_guida_2_6, linee_guida_3_1, linee_guida_4_1, linee_guida_conclusione, linee_guida_4, linee_guida_info_contatto, linee_guida_nota_motedologica, linee_guida_3_2, linee_guida_3_3, linee_guida_4_2, linee_guida_4_3, linee_guida_indice_esrs
from config import chatbot_config

########################################################################################################################

api_address = chatbot_config["api_address"]

########################################################################################################################

st.set_page_config(page_title=chatbot_config["page_title"],
                   page_icon=chatbot_config["page_icon"],
                   layout="wide",
                   initial_sidebar_state="auto",
                   menu_items=None)

if "messages" not in st.session_state:
    st.session_state.messages = copy.deepcopy(chatbot_config["messages"])

if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:6]

if "chat_history" not in st.session_state:
    st.session_state.chat_history = copy.deepcopy(chatbot_config["messages"])

if "ai_avatar_url" not in st.session_state:
    st.session_state.ai_avatar_url = chatbot_config["ai_avatar_url"]

if "user_avatar_url" not in st.session_state:
    st.session_state.user_avatar_url = chatbot_config["user_avatar_url"]

if "current_html_address" not in st.session_state:
    st.session_state.current_html_address = None
if "chat_id" not in st.session_state:
    st.session_state.chat_id = str(uuid.uuid4())[:8]  # un ID chat casuale

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

########################################################################################################################
def get_user_chats_dir(username: str) -> str:
    """
    Ritorna la directory dove salvare i file di chat per un certo utente.
    Esempio: 'username/chats/'
    """
    base_dir = username  # Puoi modificare la posizione base se vuoi
    chats_dir = os.path.join(base_dir, "chats")
    os.makedirs(chats_dir, exist_ok=True)  # Crea la directory se non esiste
    return chats_dir

def list_user_chats(username: str) -> list:
    """
    Ritorna la lista di file chat esistenti per l'utente specificato.
    Ogni file corrisponde a una chat.
    """
    chats_dir = get_user_chats_dir(username)
    all_files = os.listdir(chats_dir)
    # Filtra solo i file con estensione .json
    json_files = [f for f in all_files if f.endswith(".json")]
    # Rimuovi l'estensione .json per mostrare i nomi chat
    chat_ids = [os.path.splitext(f)[0] for f in json_files]
    return chat_ids

def load_chat_from_file(username: str, chat_id: str) -> list:
    """
    Carica la chat (lista di messaggi) dal file JSON corrispondente.
    Se il file non esiste, ritorna una lista vuota.
    """
    chats_dir = get_user_chats_dir(username)
    chat_file = os.path.join(chats_dir, f"{chat_id}.json")
    if not os.path.exists(chat_file):
        return []
    with open(chat_file, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data

def save_chat_to_file(username: str, chat_id: str, messages: list):
    """
    Salva la lista dei messaggi nel file JSON corrispondente.
    """
    chats_dir = get_user_chats_dir(username)
    chat_file = os.path.join(chats_dir, f"{chat_id}.json")
    with open(chat_file, "w", encoding="utf-8") as f:
        json.dump(messages, f, ensure_ascii=False, indent=4)
def questionnaire_page():
    """
    Funzione per mostrare la pagina del questionario.
    """
    #st.title("Questionario ESG")
    #st.markdown("Completa il questionario seguente.")

    # Renderizza il questionario e salva le risposte nella variabile di sessione
    if "questionnaire_responses" not in st.session_state:
        st.session_state["questionnaire_responses"] = {}  # Inizializza risposte vuote

    questions_file = "forms.json"
    questions = load_questions(questions_file)

    st.session_state["questionnaire_responses"] = render_questionnaire(questions)  # Usa render_quest()

    # Pulsante per salvare e tornare al chatbot
    #if st.button("Salva e torna al chatbot"):
    #    # Salva le risposte su file
    #    with open("responses.json", "w", encoding="utf-8") as f:
    #        json.dump(st.session_state["questionnaire_responses"], f, ensure_ascii=False, indent=4)

    #    st.session_state["current_page"] = "chatbot"  # Cambia pagina al chatbot
    #    st.rerun()  # Ricarica la pagina per riflettere il cambio di stato

def login_page():
    """
    Pagina di login che utilizza l'endpoint /login del backend per verificare le credenziali.
    """
    st.title("Login")

    with st.form("login_form"):
        username = st.text_input("Username", max_chars=50, placeholder="Inserisci il tuo username")
        password = st.text_input("Password", type="password", placeholder="Inserisci la tua password")
        login_button = st.form_submit_button("Login")

        if login_button:
            try:
                # Costruisci l'URL per l'endpoint di login usando api_address
                login_url = f"https://www.bluen.ai/api4/login"
                # Invia la richiesta POST con le credenziali come JSON
                response = requests.post(login_url, json={"username": username, "password": password})
                if response.status_code == 200:
                    data = response.json()
                    if data.get("login") is True:
                        st.session_state.logged_in = True
                        st.session_state["username"] = username  # <<-- AGGIUNGI QUESTA RIGA
                        st.success("Login effettuato con successo!")
                        st.rerun()  # Ricarica la pagina per mostrare il contenuto protetto
                    else:
                        st.error("Credenziali non valide.")
                else:
                    st.error(f"Errore: {response.status_code} - {response.text}")
            except Exception as e:
                st.error(f"Errore durante il login: {e}")


def documents_page():
    """
    Pagina dedicata al caricamento e alla gestione dei documenti
    """
    st.title("Gestione Documenti")

    # Sezione "Upload Documents"
    st.header("Upload Documents")
    session_id = st.session_state.session_id
    uploaded_files = st.file_uploader("Choose files", type=['pdf', 'txt', 'docx'], accept_multiple_files=True)

    if st.button("Upload and Process Documents", use_container_width=True):
        if not session_id:
            st.error("Please enter a Session ID.")
        elif not uploaded_files:
            st.error("Please upload at least one file.")
        else:
            # Process each uploaded file
            for uploaded_file in uploaded_files:
                file_id = uploaded_file.name.split(".")[0]
                description = f"Document uploaded: {file_id}"

                with st.spinner(f"Uploading and processing the document {file_id}..."):
                    # Prepare the data
                    files = {
                        'uploaded_file': (uploaded_file.name, uploaded_file.read()),
                    }
                    data = {
                        'session_id': session_id,
                        'file_id': file_id,
                        'description': description,
                    }

                    upload_document_url = f"https://www.bluen.ai/api4/upload_document"

                    try:
                        response = requests.post(upload_document_url, data=data, files=files)
                        if response.status_code == 200:
                            st.success(f"Document {file_id} uploaded and processed successfully.")
                            print(f"Upload and process response for {file_id}:", response.json())
                        else:
                            st.error(f"Failed to upload document {file_id}: {response.text}")
                            print(f"Failed to upload document {file_id}: {response.status_code}, {response.text}")
                    except Exception as e:
                        st.error(f"An error occurred: {e}")
                        print(f"An error occurred during upload for {file_id}: {e}")

            # Configure a single agent for all documents
            with st.spinner("Configuring and loading the agent for all documents..."):
                configure_chain_url = f"https://www.bluen.ai/api4/configure_and_load_chain/?session_id={session_id}"
                try:
                    response = requests.post(configure_chain_url)
                    if response.status_code == 200:
                        st.success("Agent configured and loaded successfully.")
                        print("Configure agent response:", response.json())
                    else:
                        st.error(f"Failed to configure agent: {response.text}")
                        print(f"Failed to configure agent: {response.status_code}, {response.text}")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                    print(f"An error occurred during agent configuration: {e}")

    st.markdown("---")

    # Sezione "Add URLs with Descriptions"
    st.header("Add URLs with Descriptions")

    if "url_forms" not in st.session_state:
        st.session_state.url_forms = [{"url": "", "description": ""}]

    if st.button("+ Add New URL Form", key="add_new_url_button", use_container_width=True):
        st.session_state.url_forms.append({"url": "", "description": ""})

    # Display all URL forms
    for idx, form in enumerate(st.session_state.url_forms):
        colA, colB = st.columns(2)
        with colA:
            form["url"] = st.text_input(f"URL {idx + 1}", value=form["url"], key=f"url_{idx}")
        with colB:
            form["description"] = st.text_area(f"Description {idx + 1}", value=form["description"],
                                               key=f"description_{idx}")

    if st.button("Save All URLs", key="save_all_urls_button", use_container_width=True):
        if all(form["url"] and form["description"] for form in st.session_state.url_forms):
            st.success("All URLs and descriptions saved successfully.")
        else:
            st.error("Please ensure all URL forms are filled out completely.")



def dashboard_page():
    st.title("Dashboard Admin")

    st.subheader("Crea Nuovo Utente")
    with st.form("create_user_form"):
        new_username = st.text_input("Username")
        new_password = st.text_input("Password", type="password")
        submit_new_user = st.form_submit_button("Crea Utente")
        if submit_new_user:
            register_url = f"https://www.bluen.ai/api4/register"
            # Invia la richiesta di registrazione
            response = requests.post(register_url, json={"username": new_username, "password": new_password})
            if response.status_code == 200:
                st.success("Utente creato con successo!")
            else:
                st.error(f"Errore: {response.status_code} - {response.text}")

    st.subheader("Cambia Password Utente")
    with st.form("change_password_form"):
        target_username = st.text_input("Username Utente da modificare")
        new_password_for_user = st.text_input("Nuova Password", type="password")
        submit_change = st.form_submit_button("Cambia Password")
        if submit_change:
            # Per l'admin, non serve fornire la vecchia password
            reset_url = f"https://www.bluen.ai/api4/reset_password"
            payload = {
                "requestor_username": "admin",
                "requestor_password": "admin",  # Credenziali admin hardcoded (da personalizzare se necessario)
                "target_username": target_username,
                "new_password": new_password_for_user
            }
            response = requests.post(reset_url, json=payload)
            if response.status_code == 200:
                st.success("Password cambiata con successo!")
            else:
                st.error(f"Errore: {response.status_code} - {response.text}")

    st.subheader("Lista Utenti")
    if os.path.exists("users.json"):
        with open("users.json", "r") as f:
            users = json.load(f)
        for username, hashed_password in users.items():
            # Password oscurata (ad es. mostra sempre "********")
            obfuscated = "********"
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"""
                <div style="border:1px solid #ddd; border-radius:5px; padding:10px; margin-bottom:10px;">
                    <strong>Username:</strong> {username}<br>
                    <strong>Password:</strong> {obfuscated}
                </div>
                """, unsafe_allow_html=True)
            with col2:
                if st.button("Elimina", key=f"delete_{username}"):
                    delete_url = f"{api_address}/delete_user"
                    payload = {
                        "requestor_username": "admin",
                        "requestor_password": "admin",  # Credenziali admin hardcoded
                        "target_username": username
                    }
                    response = requests.delete(delete_url, json=payload)
                    if response.status_code == 200:
                        st.success(f"Utente {username} eliminato.")
                        st.rerun()  # Ricarica la pagina per aggiornare la lista
                    else:
                        st.error(f"Errore: {response.status_code} - {response.text}")
    else:
        st.write("Nessun utente registrato.")


def download_json_file(data: dict, filename: str):
    """Converts a dictionary into a downloadable JSON file."""
    json_data = json.dumps(data, indent=4)
    b64 = base64.b64encode(json_data.encode()).decode()
    href = f'<a href="data:application/json;base64,{b64}" download="{filename}">Scarica {filename}</a>'
    return href

def is_complete_utf8(data: bytes) -> bool:
    """Verifica se `data` rappresenta una sequenza UTF-8 completa."""
    try:
        data.decode("utf-8")
        return True
    except UnicodeDecodeError:
        return False

def aggiorna_html_address(message):
    import re
    # Cerca un indirizzo che inizia con http://, contiene numeri e lettere, e termina con .html
    match = re.search(r'http://[\d\.\w\:/-]+\.html', message)
    if match:
        st.session_state.current_html_address = match.group()

    print(st.session_state.current_html_address)

def scarica_html(url):
    """Scarica il contenuto HTML da un URL."""
    try:
        risposta = requests.get(url)
        if risposta.status_code == 200:
            return risposta.text
        else:
            st.error(f"Errore nel download dell'HTML: {risposta.status_code}")
            return None
    except Exception as e:
        st.error(f"Errore durante il download dell'HTML: {e}")
        return None


def converti_html_in_docx_(html_content, output_path):
    template_path = "Carta_intestata_Bluen.docx"

    """
    Converte l'HTML in contenuto DOCX e lo aggiunge a un documento esistente con carta intestata.

    Args:
        html_content (str): Contenuto HTML da convertire.
        template_path (str): Percorso del documento DOCX esistente con carta intestata.
        output_path (str): Percorso in cui salvare il documento aggiornato.
    """
    try:
        pattern = r"<head>.*?</head>"

        # Rimuove il blocco <head> dal contenuto HTML
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)
        # Converti l'HTML in un oggetto BytesIO
        docx_io = html2docx(html_content, title="out_doc")

        # Crea un nuovo documento Word dal contenuto convertito
        temp_doc = Document(docx_io)

        # Apri il documento DOCX esistente con carta intestata
        template_doc = Document(template_path)

        # Aggiungi il contenuto HTML convertito al documento esistente
        for element in temp_doc.element.body:
            template_doc.element.body.append(element)

        # Salva il documento aggiornato con il contenuto HTML
        template_doc.save(output_path)
        print(f"Documento salvato con successo in: {output_path}")
    except Exception as e:
        print(f"Errore durante la conversione o l'inserimento: {e}")

    return output_path

def converti_html_in_docx(html_content, output_file=None):
    """
    Converte l'HTML in un file DOCX, eliminando esattamente il blocco <head>.
    """
    try:
        # Definisci una regex per catturare il blocco <head>... </head>
        pattern = r"<head>.*?</head>"

        # Rimuove il blocco <head> dal contenuto HTML
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)

        # Converti l'HTML in DOCX
        docx_file = html2docx(title="Document from HTML", content=html_content)

        # Se un file di output è specificato, salvalo
        if output_file:
            with open(output_file, "wb") as f:
                f.write(docx_file.getvalue())
            return output_file

        # Altrimenti, restituisci il contenuto DOCX come bytes
        return docx_file.getvalue()
    except Exception as e:
        st.error(f"Errore durante la conversione in DOCX: {e}")
        return None

def converti_html_in_docx__(html_content, output_path):
    template_path = "Carta_intestata_Bluen.docx"

    """
    Converte l'HTML in contenuto DOCX e lo aggiunge a un documento già esistente con carta intestata.

    Args:
        html_content (str): Contenuto HTML da convertire.
        template_path (str): Percorso del documento DOCX esistente con carta intestata.
        output_path (str): Percorso in cui salvare il documento aggiornato.
    """
    try:
        # Rimuovi il blocco <head> dal contenuto HTML
        pattern = r"<head>.*?</head>"
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)

        # Converti l'HTML in un contenuto DOCX (come bytes)
        docx_content = html2docx(title="Document from HTML", content=html_content).getvalue()

        # Apri il documento DOCX esistente con carta intestata
        template_doc = Document(template_path)

        # Crea un documento temporaneo per il contenuto HTML convertito
        temp_doc_path = "temp_doc.docx"
        with open(temp_doc_path, "wb") as temp_file:
            temp_file.write(docx_content)
        temp_doc = Document(temp_doc_path)

        # Mappa gli stili di 'temp_doc' per preservare formattazioni personalizzate
        for style in temp_doc.styles:
            if style not in template_doc.styles:
                template_doc.styles.add_style(style.name, style.type)

        # Aggiungi il contenuto HTML convertito al documento esistente
        for element in temp_doc.element.body:
            # Preserva immagini e stili personalizzati
            template_doc.element.body.append(element)

        # Salva il documento aggiornato con il contenuto HTML
        template_doc.save(output_path)
        print(f"Documento salvato con successo in: {output_path}")

        # Rimuovi il file temporaneo
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    except Exception as e:
        print(f"Errore durante la conversione o l'inserimento: {e}")

    return output_path

def download_report():
    if st.session_state.current_html_address:
        html_content = scarica_html(st.session_state.current_html_address)
        if html_content:
            docx_file = converti_html_in_docx(html_content, "report.docx")
            if docx_file:
                with open(docx_file, "rb") as f:
                    st.download_button(
                        label="Scarica il report DOCX",
                        data=f,
                        file_name="report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.error("Impossibile scaricare l'HTML dall'indirizzo fornito.")
    else:
        st.error("Nessun indirizzo HTML corrente trovato.")


def chatbot_page():

    # Sidebar for file upload and chain configuration
    #st.sidebar.header("Upload Documents")
    if st.session_state.logged_in:
        st.sidebar.subheader("Le tue chat esistenti")

        # Carica la lista delle chat per l'utente
        user_chats = list_user_chats(st.session_state["username"]) if "username" in st.session_state else []

        for c_id in user_chats:
            # Mostra ogni chat come un "bottone" o "card"
            if st.sidebar.button(f"Chat: {c_id}", key=f"load_chat_{c_id}"):
                st.session_state.chat_id = c_id
                # Carica i messaggi dal file
                loaded_msgs = load_chat_from_file(st.session_state["username"], c_id)
                st.session_state.messages = loaded_msgs
                st.session_state.chat_history = copy.deepcopy(loaded_msgs)
                st.rerun()

        # Pulsante per creare una "Nuova Chat"
        if st.sidebar.button("Nuova Chat", key="new_chat_button"):
            new_id = str(uuid.uuid4())[:8]
            st.session_state.chat_id = new_id
            st.session_state.messages = copy.deepcopy(chatbot_config["messages"])
            st.session_state.chat_history = copy.deepcopy(chatbot_config["messages"])
            st.rerun()
    # Sidebar: Add "Genera Report" button at the top
    #st.sidebar.header("Actions")
    #session_id = st.sidebar.text_input("Session ID", value="", placeholder="Inserisci Session ID")
    session_id = st.session_state.session_id
    # Sidebar for file upload and chain configuration
    """st.sidebar.header("Upload Documents")
    #session_id = st.sidebar.text_input("Session ID")
    uploaded_files = st.sidebar.file_uploader("Choose files", type=['pdf', 'txt', 'docx'], accept_multiple_files=True)

    if st.sidebar.button("Upload and Process Documents", use_container_width=True):
        if not session_id:
            st.sidebar.error("Please enter a Session ID.")
        elif not uploaded_files:
            st.sidebar.error("Please upload at least one file.")
        else:
            # Process each uploaded file
            for uploaded_file in uploaded_files:
                file_id = uploaded_file.name.split(".")[0]
                description = f"Document uploaded: {file_id}"

                with st.spinner(f"Uploading and processing the document {file_id}..."):
                    # Prepare the data
                    files = {
                        'uploaded_file': (uploaded_file.name, uploaded_file.read()),
                    }
                    data = {
                        'session_id': session_id,
                        'file_id': file_id,
                        'description': description,
                    }

                    upload_document_url = f"https://www.bluen.ai/api4/upload_document"

                    try:
                        response = requests.post(upload_document_url, data=data, files=files)
                        if response.status_code == 200:
                            st.sidebar.success(f"Document {file_id} uploaded and processed successfully.")
                            print(f"Upload and process response for {file_id}:", response.json())
                        else:
                            st.sidebar.error(f"Failed to upload document {file_id}: {response.text}")
                            print(f"Failed to upload document {file_id}: {response.status_code}, {response.text}")
                    except Exception as e:
                        st.sidebar.error(f"An error occurred: {e}")
                        print(f"An error occurred during upload for {file_id}: {e}")

            # Configure a single agent for all documents
            with st.spinner("Configuring and loading the agent for all documents..."):
                configure_chain_url = f"https://www.bluen.ai/api4/configure_and_load_chain/?session_id={session_id}"
                try:
                    response = requests.post(configure_chain_url)
                    if response.status_code == 200:
                        st.sidebar.success("Agent configured and loaded successfully.")
                        print("Configure agent response:", response.json())
                    else:
                        st.sidebar.error(f"Failed to configure agent: {response.text}")
                        print(f"Failed to configure agent: {response.status_code}, {response.text}")
                except Exception as e:
                    st.sidebar.error(f"An error occurred: {e}")
                    print(f"An error occurred during agent configuration: {e}")

    st.sidebar.markdown("---")

    # Section for adding multiple URLs with descriptions
    st.sidebar.header("Add URLs with Descriptions")

    if "url_forms" not in st.session_state:
        st.session_state.url_forms = [{"url": "", "description": ""}]

    if st.sidebar.button("+ Add New URL Form", use_container_width=True):
        st.session_state.url_forms.append({"url": "", "description": ""})

    # Display all URL forms
    for idx, form in enumerate(st.session_state.url_forms):
        #st.sidebar.markdown(f"#### URL Form {idx + 1}")
        form["url"] = st.sidebar.text_input(f"URL {idx + 1}", value=form["url"], key=f"url_{idx}")
        form["description"] = st.sidebar.text_area(f"Description {idx + 1}", value=form["description"],
                                                   key=f"description_{idx}")

    # Save all URLs at once with a single button
    if st.sidebar.button("Save All URLs", use_container_width=True):
        if all(form["url"] and form["description"] for form in st.session_state.url_forms):
            st.sidebar.success("All URLs and descriptions saved successfully.")
        else:
            st.sidebar.error("Please ensure all URL forms are filled out completely.")
            """
    st.sidebar.markdown("---")



    # Main chat interface
    previus_type = None
    for message in st.session_state.messages:

        if message["role"] == "user" and previus_type != "user":
            with st.chat_message(message["role"], avatar=st.session_state.user_avatar_url):
                st.markdown(message["content"])
        elif message["role"] != "user":
            with st.chat_message(message["role"], avatar=st.session_state.ai_avatar_url):
                st.markdown(message["content"])
        previus_type = message["role"]
    if prompt := st.chat_input("Scrivi qualcosa"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state.chat_history.append({"role": "user", "content": prompt})

        generate_response(prompt, session_id)

    st.sidebar.header("Report")

    if st.sidebar.button("Download Report", use_container_width=True):
        download_report()

    if st.sidebar.button("Genera Report", use_container_width=True):
        if session_id:
            st.sidebar.success("Generazione report in corso...")
            generate_report(session_id)
        else:
            st.sidebar.error("Per favore, inserisci un Session ID valido.")


    st.sidebar.markdown(
        """
        <div style="text-align: center; margin-top: 20px; font-size: 12px; color: black;">
            &copy; 2024 Bluen s.r.l. Tutti i diritti riservati.
        </div>
        """,
        unsafe_allow_html=True
    )

def generate_response(prompt, session_id, auto_generated=False):
    """Handles generating a response based on a prompt."""
    chain_id = f"{session_id}-workflow_generation_chain"
    input_suffix = f"""*NOTE IMPORTANTI:* 
    - Preleva informazioni utili dai vector stores, dai file in locale e dai seguenti URLs se necessario: {json.dumps(st.session_state.url_forms, indent=2) if st.session_state.url_forms else "[urls assenti]"}
    - non devi necessariamente fare grafici in qualuqnue circostanza, evita i grafici inutili e poco professionali, cioè che potrebberoe ssere trnaquillamente spiegati a parole o in tabella. focalizzati sui grafici richeisti.
    - quando generi grafici stai molto attentoa  usare etichette sintetiche per evitare che si sovrapponghino, oppure usa ad esempio colori con leggenda nei grafici a barre e altri tipi di grafici in cui nomi lunghi delle etichette deid ati su gli assi potrebbero sovrapporsi."""

    # Show the user message in the chat only if it's not auto-generated
    if not auto_generated:
        with st.chat_message("user", avatar=st.session_state.user_avatar_url):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

    if auto_generated:
        st.session_state.chat_history.append({"role": "user", "content": prompt})

    with st.chat_message("assistant", avatar=st.session_state.ai_avatar_url):
        message_placeholder = st.empty()
        s = requests.Session()
        full_response = ""
        url = f"{api_address}/chains/stream_events_chain"
        payload = {
            "chain_id": chain_id,
            "query": {
                "input": f"{prompt}\n{input_suffix}",
                "chat_history": st.session_state.chat_history[-6:] if len(st.session_state.chat_history) > 6
                else st.session_state.chat_history
            },
            "inference_kwargs": {},
        }

        non_decoded_chunk = b''
        with s.post(url, json=payload, headers=None, stream=True) as resp:
            for chunk in resp.iter_content():
                if chunk:
                    non_decoded_chunk += chunk
                    if is_complete_utf8(non_decoded_chunk):
                        decoded_chunk = non_decoded_chunk.decode("utf-8")
                        full_response += decoded_chunk
                        message_placeholder.markdown(full_response + "▌", unsafe_allow_html=True)
                        non_decoded_chunk = b''  # Clear buffer

        message_placeholder.markdown(full_response)

    # Add assistant's response to chat history if not auto-generated
    st.session_state.messages.append({"role": "assistant", "content": full_response})
    st.session_state.chat_history.append({"role": "assistant", "content": full_response})

    if "username" in st.session_state and "chat_id" in st.session_state:
        save_chat_to_file(
            st.session_state["username"],
            st.session_state["chat_id"],
            st.session_state.messages
        )

    aggiorna_html_address(full_response)
def generate_report(session_id):
    """Generates a report with predefined messages."""
    report_messages = [
    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Evita termini come 'benvenuti al' o 'nel nostro dna' etc..  Sviluppa contenuto sezione 'introduzione' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_intro}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'introduzione' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave 'Introduzione' e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 1.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_1}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli! Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli! Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 1.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 1.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_3}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 1.4 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_4}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.4 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 1.5 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_5}. {graph_notes_1_5}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.5 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto 'Descrizione Cap. 2' (nel prossimo messaggio ti chieiderò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2}. {graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Descrizione Cap. 2' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 2.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede .{linee_guida_2_1} .{graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 2.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_2}.{graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 2.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_3}.{graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 2.4 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_4}.{graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.4 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 2.5 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_5}.{graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.5 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 2.6 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_6}.{graph_notes_2}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.6 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto 'Descrizione Cap. 3' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Descrizione Cap. 3' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 3.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3_1}.{graph_notes_3_1}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 3.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 3.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3_2}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 3.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 3.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3_3}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 3.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto 'Descrizione Cap. 4' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_4}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Descrizione Cap. 4' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 4.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede.{linee_guida_4_1}. {graph_notes_4_1}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 4.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 4.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_4_2}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 4.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 4.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_4_3}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 4.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!   contenuto sezione 'Conclusione' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_conclusione}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Conclusione' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 'Nota metodologica' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_nota_motedologica}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Nota metodologica' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 'Indice ESRS' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_indice_esrs}. {graph_notes_ESRS_index}",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Indice ESRS' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Sviluppa contenuto sezione 'Informazioni di contatto' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_info_contatto}.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni che abbiano un aspetto serio e un senso di essere mostrate usando grafici complessi.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores' (non fare riferimento a vector store)... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli!  Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Informazioni di contatto' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",
]



    for message in report_messages:
        # Use auto_generated=True to avoid duplicating messages in chat history
        generate_response(message, session_id, auto_generated=True)

########################################################################################################################

def main():
    """
    Gestisce la logica principale dell'app Streamlit.
    """
    # Sidebar per la navigazione tra Chatbot e Questionario

    #col1, col2 = st.sidebar.columns(2)

    st.sidebar.markdown(
        """
        <div style="text-align: center;">
            <img src="https://static.wixstatic.com/media/63b1fb_6c4848f63b21475db8775af35ee50e55~mv2.png" alt="Logo" style="width: 300px;">
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.markdown("---")

    st.sidebar.header("Pages")

    with st.sidebar:
        if st.button("Bluen AI", key="chatbot_button", use_container_width=True):
            st.session_state["current_page"] = "chatbot"

        #with col2:
        if st.button("ESG Assessment Form", key="questionnaire_button", use_container_width=True):
            st.session_state["current_page"] = "questionnaire"

        # AGGIUNTA: bottone per la pagina Scope 2
        if st.sidebar.button("GHG Form (Scope 1)", key="scope1_button", use_container_width=True):
            st.session_state["current_page"] = "scope1"

        # AGGIUNTA: bottone per la pagina Scope 2
        if st.sidebar.button("GHG Form (Scope 2)", key="scope2_button", use_container_width=True):
            st.session_state["current_page"] = "scope2"

        # AGGIUNTA: bottone per la pagina Scope 2
        if st.sidebar.button("GHG Form (Scope 3)", key="scope3_button", use_container_width=True):
            st.session_state["current_page"] = "scope3"

        if st.button("Documents", key="documents_button", use_container_width=True):
            st.session_state["current_page"] = "documents"

        # Se l'utente loggato è admin, mostra il bottone per la Dashboard
        if st.session_state.get("username") == "admin":
            if st.sidebar.button("Dashboard", key="dashboard_button", use_container_width=True):
                st.session_state["current_page"] = "dashboard"

    st.sidebar.markdown("---")

    if st.session_state.get("current_page", "chatbot") == "chatbot":
        chatbot_page()  # Funzione che gestisce la logica del chatbot
    elif st.session_state.get("current_page") == "questionnaire":
        questionnaire_page()
    elif st.session_state.get("current_page") == "scope1":
        render_scope1_form(session_id=st.session_state.session_id)
    elif st.session_state.get("current_page") == "scope2":
        render_scope2_form(session_id=st.session_state.session_id)
    elif st.session_state.get("current_page") == "scope3":
        render_scope3_form(session_id=st.session_state.session_id)
    elif st.session_state.get("current_page") == "dashboard":  # <<-- AGGIUNGI QUESTA CONDIZIONE
        dashboard_page()
    elif st.session_state.get("current_page") == "documents":
        documents_page()

# Se l'utente non è loggato, mostra la login_page
if not st.session_state.logged_in:
    login_page()
else:
    # Se loggato, richiama main()
    main()
