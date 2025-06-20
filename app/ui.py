import json
import os
import base64
import uuid
from typing import Any
import requests
import streamlit as st
import copy
import re
from docx import Document
from html2docx import html2docx
import os

from prompt_utils import gather_all_forms_data
#-
from scope3_form import render_scope3_form
from scope1_form import render_scope1_form
from scope2_form import render_scope2_form
from forms_ui import render_questionnaire, load_questions, merge_questions_with_responses
#from utilities import graph_notes_2, graph_notes_1_5, graph_notes_3_1, graph_notes_4_1, graph_notes_ESRS_index
from utilities import linee_guida_intro, linee_guida_1_1, linee_guida_1_2, linee_guida_1_3, linee_guida_2, linee_guida_3, linee_guida_1_4, linee_guida_1_5, linee_guida_2_1, linee_guida_2_2, linee_guida_2_3, linee_guida_2_4, linee_guida_2_5, linee_guida_2_6, linee_guida_3_1, linee_guida_4_1, linee_guida_conclusione, linee_guida_4, linee_guida_info_contatto, linee_guida_nota_motedologica, linee_guida_3_2, linee_guida_3_3, linee_guida_4_2, linee_guida_4_3, linee_guida_indice_esrs
from config import chatbot_config

########################################################################################################################

api_address = chatbot_config["api_address"]

########################################################################################################################

st.set_page_config(page_title=chatbot_config["page_title"],
                   page_icon=chatbot_config["page_icon"],
                   layout="wide",
                   initial_sidebar_state="auto",
                   menu_items=None)

if "messages" not in st.session_state:
    st.session_state.messages = copy.deepcopy(chatbot_config["messages"])

if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:6]

if "chat_history" not in st.session_state:
    st.session_state.chat_history = copy.deepcopy(chatbot_config["messages"])

if "ai_avatar_url" not in st.session_state:
    st.session_state.ai_avatar_url = chatbot_config["ai_avatar_url"]

if "user_avatar_url" not in st.session_state:
    st.session_state.user_avatar_url = chatbot_config["user_avatar_url"]

if "current_html_address" not in st.session_state:
    st.session_state.current_html_address = None
if "chat_id" not in st.session_state:
    st.session_state.chat_id = str(uuid.uuid4())[:8]  # un ID chat casuale

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

########################################################################################################################

def load_forms_and_urls_data():
    username = st.session_state.get("username")
    session_id = st.session_state.get("session_id", "default_session")

    # Caricamento del questionario
    if username:
        forms_dir = get_user_session_forms_dir(username, session_id)
        q_file_path = os.path.join(forms_dir, "questionnaire.json")
        if os.path.exists(q_file_path):
            with open(q_file_path, "r", encoding="utf-8") as f:
                saved_responses = json.load(f)
        else:
            saved_responses = {}

        st.session_state["questionnaire_responses"] = saved_responses
        # Aggiorna inoltre la versione completa se esiste la funzione merge
        questions = load_questions()  # Presupponendo che questa funzione esista
        st.session_state["full_questionnaire_responses"] = merge_questions_with_responses(questions, saved_responses)
    else:
        st.warning("Non sei loggato, i dati del questionario non potranno essere salvati/caricati.")

    # Caricamento degli URL
    if username:
        urls_dir = get_user_session_urls_dir(username, session_id)
        urls_file_path = os.path.join(urls_dir, "urls_list.json")
        if os.path.exists(urls_file_path):
            with open(urls_file_path, "r", encoding="utf-8") as f:
                saved_urls = json.load(f)
            st.session_state["url_forms"] = saved_urls
        else:
            st.session_state["url_forms"] = []
    else:
        st.warning("Non sei loggato, non posso caricare gli URL salvati.")


def list_user_sessions(username: str) -> list:
    """
    Restituisce la lista di 'session_id' già esistenti per un dato utente,
    andando a verificare le sottocartelle in USERS_DATA/<username>/.
    """
    base_path = os.path.join("USERS_DATA", username)
    if not os.path.isdir(base_path):
        return []

    # Filtra solo le sottocartelle (ognuna dovrebbe essere un session_id)
    possible_sessions = []
    for item in os.listdir(base_path):
        full_item_path = os.path.join(base_path, item)
        if os.path.isdir(full_item_path):
            # item potrebbe essere un session_id
            possible_sessions.append(item)

    return sorted(possible_sessions)


def get_user_session_dir(username: str, session_id: str) -> str:
    """
    Ritorna il percorso della cartella dedicata all'utente e alla specifica session_id.
    Esempio: 'mario/AB12/'
    """
    user_dir = username
    session_dir = os.path.join("USERS_DATA", user_dir, session_id)
    os.makedirs(session_dir, exist_ok=True)
    return session_dir

def get_user_session_files_dir(username: str, session_id: str) -> str:
    """
    Ritorna il percorso in cui salvare i file caricati (PDF, docx, etc.)
    Esempio: 'mario/AB12/files/'
    """
    base_dir = get_user_session_dir(username, session_id)
    files_dir = os.path.join(base_dir, "files")
    os.makedirs(files_dir, exist_ok=True)
    return files_dir

def get_user_session_urls_dir(username: str, session_id: str) -> str:
    """
    Ritorna il percorso in cui salvare le risorse delle URL (es. HTML scaricato, etc.).
    Esempio: 'mario/AB12/urls/'
    """
    base_dir = get_user_session_dir(username, session_id)
    urls_dir = os.path.join(base_dir, "urls")
    os.makedirs(urls_dir, exist_ok=True)
    return urls_dir

def get_user_session_forms_dir(username: str, session_id: str) -> str:
    """
    Ritorna il percorso in cui salvare i JSON dei forms.
    Esempio: 'mario/AB12/forms/'
    """
    base_dir = get_user_session_dir(username, session_id)
    forms_dir = os.path.join(base_dir, "forms")
    os.makedirs(forms_dir, exist_ok=True)
    return forms_dir

def get_user_reports_dir(username: str) -> str:
    """
    Ritorna la directory dove salvare i file di report per un certo utente.
    Esempio: 'username/reports/'
    """
    base_dir = username  # cartella base per l'utente
    reports_dir = os.path.join(base_dir, "reports")
    os.makedirs(reports_dir, exist_ok=True)  # Crea la directory se non esiste
    return reports_dir

def get_user_chats_dir(username: str) -> str:
    """
    Ritorna la directory dove salvare i file di chat per un certo utente.
    Esempio: 'username/chats/'
    """
    base_dir = username  # Puoi modificare la posizione base se vuoi
    chats_dir = os.path.join(base_dir, "chats")
    os.makedirs(chats_dir, exist_ok=True)  # Crea la directory se non esiste
    return chats_dir

def list_user_chats(username: str) -> list:
    """
    Ritorna la lista di file chat esistenti per l'utente specificato.
    Ogni file corrisponde a una chat.
    """
    chats_dir = get_user_chats_dir(username)
    all_files = os.listdir(chats_dir)
    # Filtra solo i file con estensione .json
    json_files = [f for f in all_files if f.endswith(".json")]
    # Rimuovi l'estensione .json per mostrare i nomi chat
    chat_ids = [os.path.splitext(f)[0] for f in json_files]
    return chat_ids

def load_chat_from_file(username: str, chat_id: str) -> dict:
    """
    Carica la chat (dizionario) dal file JSON corrispondente.
    Se il file non esiste, ritorna un dizionario con nome e messages vuoti.
    """
    chats_dir = get_user_chats_dir(username)
    chat_file = os.path.join(chats_dir, f"{chat_id}.json")
    if not os.path.exists(chat_file):
        return {
            "id": chat_id,
            "name": "Nuova Chat",
            "messages": []
        }
    with open(chat_file, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data

def save_chat_to_file(username: str, chat_id: str, chat_data: dict):
    """
    Salva la struttura della chat (id, name, messages) nel file JSON corrispondente.
    """
    chats_dir = get_user_chats_dir(username)
    chat_file = os.path.join(chats_dir, f"{chat_id}.json")
    with open(chat_file, "w", encoding="utf-8") as f:
        json.dump(chat_data, f, ensure_ascii=False, indent=4)


def save_form_responses_locally(username: str, session_id: str, form_data: dict, filename: str = "questionnaire.json"):
    forms_dir = get_user_session_forms_dir(username, session_id)
    file_path = os.path.join(forms_dir, filename)
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(form_data, f, ensure_ascii=False, indent=4)


def questionnaire_page():
    """
    Pagina Streamlit per il questionario ESG con supporto per il salvataggio e il recupero delle risposte salvate.
    """

    username = st.session_state.get("username")
    session_id = st.session_state.get("session_id", "default_session")

    st.title("Questionario ESG")

    # Bottone di salvataggio locale
    if st.button("Salva questionario in locale", use_container_width=True):
        if username:
            save_form_responses_locally(username, session_id,
                                        st.session_state["questionnaire_responses"],
                                        filename="questionnaire.json")
            st.success("Questionario salvato localmente!")
        else:
            st.warning("Devi essere loggato per salvare localmente.")

    # 1) Carica risposte se esiste forms/questionnaire.json
    saved_responses = {}  # Dizionario vuoto per evitare errori se non esiste il file

    if username:
        forms_dir = get_user_session_forms_dir(username, session_id)
        q_file_path = os.path.join(forms_dir, "questionnaire.json")

        if os.path.exists(q_file_path):
            with open(q_file_path, "r", encoding="utf-8") as f:
                saved_responses = json.load(f)

        # Memorizziamo in session_state per utilizzo futuro
        st.session_state["questionnaire_responses"] = saved_responses
    else:
        st.warning("Non sei loggato, i dati del questionario non potranno essere salvati/caricati.")

    # 2) Carichiamo le domande
    questions_file = "forms.json"
    questions = load_questions()
    # 3) Render del questionario, passando le risposte salvate
    st.session_state["questionnaire_responses"] = render_questionnaire(questions, saved_responses)
    st.session_state["full_questionnaire_responses"] = merge_questions_with_responses(questions, saved_responses)
    print(st.session_state["full_questionnaire_responses"])


def login_page():
    """
    Pagina di login che utilizza l'endpoint /login del backend per verificare le credenziali.
    """
    st.title("Login")

    with st.form("login_form"):
        username = st.text_input("Username", max_chars=50, placeholder="Inserisci il tuo username")
        password = st.text_input("Password", type="password", placeholder="Inserisci la tua password")
        login_button = st.form_submit_button("Login")

        if login_button:
            try:
                # Costruisci l'URL per l'endpoint di login usando api_address
                login_url = f"http://34.91.209.79:8000/login"
                # Invia la richiesta POST con le credenziali come JSON
                response = requests.post(login_url, json={"username": username, "password": password})
                if response.status_code == 200:
                    data = response.json()
                    if data.get("login") is True:
                        st.session_state.logged_in = True
                        st.session_state["username"] = username  # <<-- AGGIUNGI QUESTA RIGA
                        st.success("Login effettuato con successo!")
                        st.rerun()  # Ricarica la pagina per mostrare il contenuto protetto
                    else:
                        st.error("Credenziali non valide.")
                else:
                    st.error(f"Errore: {response.status_code} - {response.text}")
            except Exception as e:
                st.error(f"Errore durante il login: {e}")


def documents_page():
    """
    Pagina dedicata al caricamento e alla gestione dei documenti
    """
    import os
    import json
    import base64
    import uuid
    from pathlib import Path

    st.title("Gestione Documenti")

    # Recupera username e session_id dalla sessione
    username = st.session_state.get("username", None)
    session_id = st.session_state.session_id

    # --- Nuovo: Pulsante per configurare la catena ---
    st.header("Configurazione della Catena")
    if st.button("Configura e Carica Catena", key="configure_chain_button", use_container_width=True):
        with st.spinner("Configurazione e caricamento della catena in corso..."):
            configure_chain_url = f"http://34.91.209.79:8000/configure_and_load_chain/?session_id={session_id}"
            try:
                response = requests.post(configure_chain_url)
                if response.status_code == 200:
                    st.success("Catena configurata e caricata con successo.")
                    print("Configure agent response:", response.json())
                else:
                    st.error(f"Errore durante la configurazione della catena: {response.text}")
                    print(f"Failed to configure agent: {response.status_code}, {response.text}")
            except Exception as e:
                st.error(f"Errore durante la configurazione della catena: {e}")
                print(f"An error occurred during agent configuration: {e}")

    # 1) Chiave del widget file_uploader (per poterlo ricreare e azzerare)
    # Se non esiste ancora in session_state, la creiamo
    if "file_upload_widget_key" not in st.session_state:
        st.session_state.file_upload_widget_key = "file_upload_main"

    # --------------------------------------- #
    # SEZIONE: Upload Documents
    # --------------------------------------- #
    st.header("Upload Documents")

    # Widget file_uploader con la chiave (variabile) in session_state
    uploaded_files = st.file_uploader(
        "Choose files",
        type=['pdf', 'txt', 'docx'],
        accept_multiple_files=True,
        key=st.session_state.file_upload_widget_key
    )

    if st.button("Upload and Process Documents", use_container_width=True):
        if not session_id:
            st.error("Please enter a Session ID.")
        elif not uploaded_files:
            st.error("Please upload at least one file.")
        else:
            # Process each uploaded file
            for uploaded_file in uploaded_files:
                file_id = uploaded_file.name.split(".")[0]
                description = f"Document uploaded: {file_id}"

                with st.spinner(f"Uploading and processing the document {file_id}..."):
                    # 1) Leggi il contenuto del file in memoria
                    file_bytes = uploaded_file.read()

                    # 2) Se l'utente è loggato, salviamo localmente
                    if username:
                        local_files_dir = get_user_session_files_dir(username, session_id)
                        local_file_path = os.path.join(local_files_dir, uploaded_file.name)

                        # Salva fisicamente il file in locale
                        with open(local_file_path, "wb") as f:
                            f.write(file_bytes)
                    else:
                        local_file_path = None
                        st.warning("Non sei loggato, il file non verrà salvato in locale.")

                    # 3) Preparazione parametri POST
                    data = {
                        'session_id': session_id,
                        'file_id': file_id,
                        'description': description,
                    }
                    upload_document_url = f"http://34.91.209.79:8000/upload_document"

                    # 4) Prepara il file per la POST
                    if local_file_path and os.path.exists(local_file_path):
                        # Riapriamo il file salvato in locale
                        files = {
                            'uploaded_file': (uploaded_file.name, open(local_file_path, "rb")),
                        }
                    else:
                        # Altrimenti usiamo i bytes in memoria
                        files = {
                            'uploaded_file': (uploaded_file.name, file_bytes),
                        }

                    # 5) Invio la richiesta al backend
                    try:
                        response = requests.post(upload_document_url, data=data, files=files)
                        if response.status_code == 200:
                            st.success(f"Document {file_id} uploaded and processed successfully.")
                            print(f"Upload and process response for {file_id}:", response.json())
                        else:
                            st.error(f"Failed to upload document {file_id}: {response.text}")
                            print(f"Failed to upload document {file_id}: {response.status_code}, {response.text}")
                    except Exception as e:
                        st.error(f"An error occurred: {e}")
                        print(f"An error occurred during upload for {file_id}: {e}")

            # 6) Configurazione e caricamento della chain
            #with st.spinner("Configuring and loading the agent for all documents..."):
                # TODO:
                #  - aggiungi eliminazione catena e scaricamento da memoria, cosi da ricrearla ogni volta
            #    configure_chain_url = f"http://34.91.209.79:8000/configure_and_load_chain/?session_id={session_id}"
            #    try:
            #        response = requests.post(configure_chain_url)
            #        if response.status_code == 200:
            #            st.success("Agent configured and loaded successfully.")
            #            print("Configure agent response:", response.json())
            #        else:
            #            st.error(f"Failed to configure agent: {response.text}")
            #            print(f"Failed to configure agent: {response.status_code}, {response.text}")
            #    except Exception as e:
            #        st.error(f"An error occurred: {e}")
            #        print(f"An error occurred during agent configuration: {e}")

        # Al termine di TUTTI gli upload, cambiamo la chiave del widget → svuotiamo la lista
        st.session_state.file_upload_widget_key = "file_upload_" + str(uuid.uuid4())[:6]
        st.rerun()

    # --------------------------------------- #
    # MOSTRA FILE CARICATI LOCALMENTE (con card + link download)
    # --------------------------------------- #
    st.markdown("---")
    if username:
        files_dir = get_user_session_files_dir(username, session_id)
        st.subheader("File caricati localmente (per questa sessione)")

        if os.path.exists(files_dir):
            local_files = sorted(Path(files_dir).glob("*.*"))
            if local_files:
                # Contenitore scrollabile
                st.markdown("<div style='max-height:300px; overflow-y:auto;'>", unsafe_allow_html=True)

                for local_file_path in local_files:
                    file_name = local_file_path.name

                    # Creiamo un link testuale cliccabile per il download
                    with open(local_file_path, "rb") as f:
                        file_data = f.read()
                    b64_file = base64.b64encode(file_data).decode()
                    download_link = f'<a href="data:application/octet-stream;base64,{b64_file}" download="{file_name}">Clicca per scaricare</a>'

                    st.markdown(f"""
                    <div style="border:1px solid #ddd; border-radius:5px; padding:10px; margin-bottom:10px;">
                        <strong>{file_name}</strong><br>
                        {download_link}
                    </div>
                    """, unsafe_allow_html=True)

                st.markdown("</div>", unsafe_allow_html=True)
            else:
                st.info("Nessun file presente nella cartella locale.")
        else:
            st.info("La cartella files per questa sessione non esiste ancora.")
    else:
        st.warning("Effettua il login per visualizzare i file caricati localmente.")

    # --------------------------------------- #
    # SEZIONE: Add URLs with Descriptions
    # --------------------------------------- #
    st.markdown("---")
    st.header("Add URLs with Descriptions")

    # Per la singola form di URL
    if "new_url_value" not in st.session_state:
        st.session_state.new_url_value = ""

    if "new_url_desc" not in st.session_state:
        st.session_state.new_url_desc = ""

    # Input e text_area unici
    st.session_state.new_url_value = st.text_input("URL", value=st.session_state.new_url_value, key="url_unique")
    st.session_state.new_url_desc = st.text_area("Description", value=st.session_state.new_url_desc, key="desc_unique")

    # Pulsante di salvataggio
    if st.button("Salva URL", key="save_single_url_button", use_container_width=True):
        # Verifichiamo che non siano vuoti
        if st.session_state.new_url_value.strip() and st.session_state.new_url_desc.strip():
            # Salviamo nel file
            if username:
                urls_dir = get_user_session_urls_dir(username, session_id)
                urls_file_path = os.path.join(urls_dir, "urls_list.json")

                # Leggiamo la lista esistente (se c'è)
                if os.path.exists(urls_file_path):
                    with open(urls_file_path, "r", encoding="utf-8") as f:
                        saved_urls = json.load(f)
                else:
                    saved_urls = []

                # Append la nuova URL
                saved_urls.append({
                    "url": st.session_state.new_url_value,
                    "description": st.session_state.new_url_desc
                })

                # Riscriviamo
                with open(urls_file_path, "w", encoding="utf-8") as f:
                    json.dump(saved_urls, f, ensure_ascii=False, indent=4)

                st.success("URL salvata con successo!")

                # Reset campi input
                st.session_state.new_url_value = ""
                st.session_state.new_url_desc = ""

                st.session_state["url_forms"] = saved_urls

                st.rerun()
            else:
                st.warning("Devi essere loggato per salvare localmente.")
        else:
            st.warning("Compila sia URL che Description.")

    # MOSTRA GLI URL SALVATI LOCALMENTE (card style, in contenitore scrollabile)
    if username:
        urls_dir = get_user_session_urls_dir(username, session_id)
        urls_file_path = os.path.join(urls_dir, "urls_list.json")
        if os.path.exists(urls_file_path):
            with open(urls_file_path, "r", encoding="utf-8") as f:
                saved_urls = json.load(f)

            st.subheader("URL salvati localmente (per questa sessione)")
            if saved_urls:
                # Contenitore scrollabile
                st.markdown("<div style='max-height:300px; overflow-y:auto;'>", unsafe_allow_html=True)

                for i, item in enumerate(saved_urls, start=1):
                    url_val = item["url"]
                    desc = item["description"]
                    st.markdown(f"""
                    <div style="border:1px solid #ddd; border-radius:5px; padding:10px; margin-bottom:10px;">
                        <strong>URL #{i}</strong><br>
                        <em>Link:</em> <a href="{url_val}" target="_blank">{url_val}</a><br>
                        <em>Descrizione:</em> {desc}
                    </div>
                    """, unsafe_allow_html=True)

                st.markdown("</div>", unsafe_allow_html=True)
            else:
                st.info("Non ci sono URL salvati.")
        else:
            st.info("Nessun file urls_list.json trovato nella cartella locale.")
    else:
        st.warning("Effettua il login per visualizzare gli URL salvati.")


def reports_page():
    st.title("I tuoi Report")

    # Recupera il session_id dalla sessione
    session_id = st.session_state.get("session_id")
    if not session_id:
        st.error("Session ID non disponibile.")
        return

    # Costruisci l'URL del file HTML remoto
    url = f"http://34.91.209.79:8093/files/{session_id}/bilancio_sostenibilita.html"

    try:
        response = requests.get(url)
        if response.status_code == 200:
            html_content = response.text
            # Visualizza il contenuto HTML renderizzato nella pagina
            st.components.v1.html(html_content, height=800, scrolling=True)
        else:
            st.warning("Il file HTML non è presente.")
    except Exception as e:
        st.error("Errore durante il caricamento del file HTML.")


########################################
# Funzioni di utilità per API Keys
########################################

def load_api_keys(filename="api_keys.json"):
    """
    Carica l'elenco delle chiavi API da un file JSON.
    Se il file non esiste o è vuoto, restituisce una lista vuota.
    """
    if not os.path.exists(filename):
        return []

    with open(filename, "r", encoding="utf-8") as f:
        try:
            data = json.load(f)
            return data.get("openai_api_keys", [])
        except json.JSONDecodeError:
            return []

def save_api_keys(api_keys, filename="api_keys.json"):
    """
    Salva l'elenco delle chiavi API in un file JSON.
    """
    data = {"openai_api_keys": api_keys}
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def generate_key_id():
    """
    Genera un ID univoco (brevi) per identificare ogni chiave.
    """
    return str(uuid.uuid4())[:8]

########################################
# Esempio di modifica al dashboard_page
########################################

def dashboard_page():
    st.title("Dashboard Admin")

    # ---------------------------- #
    # Sezione: Creazione Nuovo Utente
    # ---------------------------- #
    st.subheader("Crea Nuovo Utente")
    with st.form("create_user_form"):
        new_username = st.text_input("Username")
        new_password = st.text_input("Password", type="password")
        submit_new_user = st.form_submit_button("Crea Utente")
        if submit_new_user:
            register_url = f"http://34.91.209.79:8000/register"
            response = requests.post(register_url, json={"username": new_username, "password": new_password})
            if response.status_code == 200:
                st.success("Utente creato con successo!")
            else:
                st.error(f"Errore: {response.status_code} - {response.text}")

    # ---------------------------- #
    # Sezione: Cambia Password Utente
    # ---------------------------- #
    st.subheader("Cambia Password Utente")
    with st.form("change_password_form"):
        target_username = st.text_input("Username Utente da modificare")
        new_password_for_user = st.text_input("Nuova Password", type="password")
        submit_change = st.form_submit_button("Cambia Password")
        if submit_change:
            reset_url = f"{api_address}/reset_password"
            payload = {
                "requestor_username": "admin",
                "requestor_password": "admin",  # Credenziali admin hardcoded
                "target_username": target_username,
                "new_password": new_password_for_user
            }
            response = requests.post(reset_url, json=payload)
            if response.status_code == 200:
                st.success("Password cambiata con successo!")
            else:
                st.error(f"Errore: {response.status_code} - {response.text}")

    st.subheader("Lista Utenti")
    if os.path.exists("users.json"):
        with open("users.json", "r") as f:
            users = json.load(f)
        for username, hashed_password in users.items():
            # Password oscurata
            obfuscated = "********"
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"""
                <div style="border:1px solid #ddd; border-radius:5px; padding:10px; margin-bottom:10px;">
                    <strong>Username:</strong> {username}<br>
                    <strong>Password:</strong> {obfuscated}
                </div>
                """, unsafe_allow_html=True)
            with col2:
                if st.button("Elimina", key=f"delete_{username}"):
                    delete_url = f"{api_address}/delete_user"
                    payload = {
                        "requestor_username": "admin",
                        "requestor_password": "admin",
                        "target_username": username
                    }
                    response = requests.delete(delete_url, json=payload)
                    if response.status_code == 200:
                        st.success(f"Utente {username} eliminato.")
                        st.rerun()
                    else:
                        st.error(f"Errore: {response.status_code} - {response.text}")
    else:
        st.write("Nessun utente registrato.")

    # ---------------------------- #
    # SEZIONE: Gestione Chiavi OpenAI
    # ---------------------------- #
    st.subheader("Gestione Chiavi OpenAI")

    # Carica le chiavi esistenti dal file
    api_keys = load_api_keys()

    # Inizializza (se non esiste) lo stato per mostrare/nascondere le chiavi
    if "show_api_keys" not in st.session_state:
        st.session_state["show_api_keys"] = False

    # Form per aggiungere una nuova chiave
    with st.form("add_api_key_form"):
        new_api_key = st.text_input("Inserisci nuova API Key OpenAI", type="password")
        add_key_btn = st.form_submit_button("Aggiungi Chiave")
        if add_key_btn and new_api_key.strip():
            # Crea un record con ID e chiave
            new_record = {
                "id": generate_key_id(),
                "key": new_api_key.strip()
            }
            api_keys.append(new_record)
            save_api_keys(api_keys)
            st.success("Nuova API Key aggiunta con successo!")
            st.rerun()

    # Bottone toggle: mostra/nascondi chiavi
    if st.button("Mostra/Nascondi Chiavi"):
        st.session_state["show_api_keys"] = not st.session_state["show_api_keys"]

    # Se ci sono chiavi, visualizzale (mascherate o in chiaro)
    if api_keys:
        st.markdown("#### Chiavi API esistenti:")
        for record in api_keys:
            key_id = record["id"]
            key_value = record["key"]

            colA, colB, colC = st.columns([1.2, 3, 1])
            with colA:
                st.markdown(f"**ID:** `{key_id}`")
            with colB:
                if st.session_state["show_api_keys"]:
                    # Mostra la chiave in chiaro
                    st.markdown(f"**Key:** `{key_value}`")
                else:
                    # Oscura la chiave (es. sostituisce tutti i caratteri con asterischi)
                    masked = "•" * len(key_value)
                    st.markdown(f"**Key:** `{masked}`")
            with colC:
                # Pulsante di rimozione
                if st.button("Rimuovi", key=f"remove_{key_id}"):
                    # Rimuovi la chiave dall'elenco
                    api_keys = [k for k in api_keys if k["id"] != key_id]
                    save_api_keys(api_keys)
                    st.success(f"Chiave con ID {key_id} rimossa.")
                    st.rerun()
    else:
        st.markdown("*Nessuna chiave API salvata.*")


def download_json_file(data: dict, filename: str):
    """Converts a dictionary into a downloadable JSON file."""
    json_data = json.dumps(data, indent=4)
    b64 = base64.b64encode(json_data.encode()).decode()
    href = f'<a href="data:application/json;base64,{b64}" download="{filename}">Scarica {filename}</a>'
    return href

def is_complete_utf8(data: bytes) -> bool:
    """Verifica se `data` rappresenta una sequenza UTF-8 completa."""
    try:
        data.decode("utf-8")
        return True
    except UnicodeDecodeError:
        return False

def aggiorna_html_address(message):
    import re
    # Cerca un indirizzo che inizia con http://, contiene numeri e lettere, e termina con .html
    match = re.search(r'http://[\d\.\w\:/-]+\.html', message)
    if match:
        st.session_state.current_html_address = match.group()

    print(st.session_state.current_html_address)

def scarica_html(url):
    """Scarica il contenuto HTML da un URL."""
    try:
        risposta = requests.get(url)
        if risposta.status_code == 200:
            return risposta.text
        else:
            st.error(f"Errore nel download dell'HTML: {risposta.status_code}")
            return None
    except Exception as e:
        st.error(f"Errore durante il download dell'HTML: {e}")
        return None


def converti_html_in_docx_(html_content, output_path):
    template_path = "Carta_intestata_Bluen.docx"

    """
    Converte l'HTML in contenuto DOCX e lo aggiunge a un documento esistente con carta intestata.

    Args:
        html_content (str): Contenuto HTML da convertire.
        template_path (str): Percorso del documento DOCX esistente con carta intestata.
        output_path (str): Percorso in cui salvare il documento aggiornato.
    """
    try:
        pattern = r"<head>.*?</head>"

        # Rimuove il blocco <head> dal contenuto HTML
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)
        # Converti l'HTML in un oggetto BytesIO
        docx_io = html2docx(html_content, title="out_doc")

        # Crea un nuovo documento Word dal contenuto convertito
        temp_doc = Document(docx_io)

        # Apri il documento DOCX esistente con carta intestata
        template_doc = Document(template_path)

        # Aggiungi il contenuto HTML convertito al documento esistente
        for element in temp_doc.element.body:
            template_doc.element.body.append(element)

        # Salva il documento aggiornato con il contenuto HTML
        template_doc.save(output_path)
        print(f"Documento salvato con successo in: {output_path}")
    except Exception as e:
        print(f"Errore durante la conversione o l'inserimento: {e}")

    return output_path

def converti_html_in_docx(html_content, output_file=None):
    """
    Converte l'HTML in un file DOCX, eliminando esattamente il blocco <head>.
    """
    try:
        # Definisci una regex per catturare il blocco <head>... </head>
        pattern = r"<head>.*?</head>"

        # Rimuove il blocco <head> dal contenuto HTML
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)

        # Converti l'HTML in DOCX
        docx_file = html2docx(title="Document from HTML", content=html_content)

        # Se un file di output è specificato, salvalo
        if output_file:
            with open(output_file, "wb") as f:
                f.write(docx_file.getvalue())
            return output_file

        # Altrimenti, restituisci il contenuto DOCX come bytes
        return docx_file.getvalue()
    except Exception as e:
        st.error(f"Errore durante la conversione in DOCX: {e}")
        return None

def converti_html_in_docx__(html_content, output_path):
    template_path = "Carta_intestata_Bluen.docx"

    """
    Converte l'HTML in contenuto DOCX e lo aggiunge a un documento già esistente con carta intestata.

    Args:
        html_content (str): Contenuto HTML da convertire.
        template_path (str): Percorso del documento DOCX esistente con carta intestata.
        output_path (str): Percorso in cui salvare il documento aggiornato.
    """
    try:
        # Rimuovi il blocco <head> dal contenuto HTML
        pattern = r"<head>.*?</head>"
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)

        # Converti l'HTML in un contenuto DOCX (come bytes)
        docx_content = html2docx(title="Document from HTML", content=html_content).getvalue()

        # Apri il documento DOCX esistente con carta intestata
        template_doc = Document(template_path)

        # Crea un documento temporaneo per il contenuto HTML convertito
        temp_doc_path = "temp_doc.docx"
        with open(temp_doc_path, "wb") as temp_file:
            temp_file.write(docx_content)
        temp_doc = Document(temp_doc_path)

        # Mappa gli stili di 'temp_doc' per preservare formattazioni personalizzate
        for style in temp_doc.styles:
            if style not in template_doc.styles:
                template_doc.styles.add_style(style.name, style.type)

        # Aggiungi il contenuto HTML convertito al documento esistente
        for element in temp_doc.element.body:
            # Preserva immagini e stili personalizzati
            template_doc.element.body.append(element)

        # Salva il documento aggiornato con il contenuto HTML
        template_doc.save(output_path)
        print(f"Documento salvato con successo in: {output_path}")

        # Rimuovi il file temporaneo
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    except Exception as e:
        print(f"Errore durante la conversione o l'inserimento: {e}")

    return output_path

def download_report():
    if st.session_state.current_html_address:
        html_content = scarica_html(st.session_state.current_html_address)
        if html_content:
            docx_file = converti_html_in_docx(html_content, "report.docx")
            if docx_file:
                with open(docx_file, "rb") as f:
                    st.download_button(
                        label="Scarica il report DOCX",
                        data=f,
                        file_name="report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.error("Impossibile scaricare l'HTML dall'indirizzo fornito.")
    else:
        st.error("Nessun indirizzo HTML corrente trovato.")

def workspace_management_page():
    """
    Pagina per creare, visualizzare ed eliminare gli Spazi di Lavoro.
    (A livello di codice useremo la logica session_id, ma in UI mostriamo 'Spazio di Lavoro').
    """
    st.title("Gestione Spazi di Lavoro")

    username = st.session_state.get("username", None)
    if not username:
        st.error("Devi essere loggato per gestire i tuoi Spazi di Lavoro.")
        return

    # 1) Creazione di un nuovo Spazio di Lavoro
    st.subheader("Crea un nuovo Spazio di Lavoro")
    new_workspace_name = st.text_input("Nome per il tuo nuovo Spazio di Lavoro", placeholder="es. Progetto Alpha")

    if st.button("Crea Spazio di Lavoro", use_container_width=True):
        if new_workspace_name.strip():
            # Genera un nuovo session_id (a livello di codice)
            new_session_id = str(uuid.uuid4())[:6]
            # Crea la cartella associata all'utente e a questo session_id
            created_path = get_user_session_dir(username, new_session_id)

            # Aggiungi queste RIGHE:
            workspace_name_path = os.path.join(created_path, "workspace_name.txt")
            with open(workspace_name_path, "w", encoding="utf-8") as f:
                f.write(new_workspace_name.strip())

            st.success(f"Spazio di Lavoro '{new_workspace_name}' creato con ID: {new_session_id}")
        else:
            st.warning("Inserisci un nome valido per lo Spazio di Lavoro.")

    st.markdown("---")

    # 2) Elenco Spazi di Lavoro esistenti
    st.subheader("I tuoi Spazi di Lavoro esistenti")
    existing_workspaces = list_user_sessions(username)

    if not existing_workspaces:
        st.info("Non ci sono ancora Spazi di Lavoro.")
    else:
        for sid in existing_workspaces:
            # Mostra una card con info (session_id = sid), pulsante per selezionarlo o per eliminarlo
            col1, col2, col3 = st.columns([3, 1, 1])
            with col1:
                # Aggiungi queste righe per leggere il nome se esiste:
                session_path = get_user_session_dir(username, sid)
                name_file = os.path.join(session_path, "workspace_name.txt")
                if os.path.exists(name_file):
                    with open(name_file, "r", encoding="utf-8") as f:
                        workspace_name = f.read().strip()
                else:
                    workspace_name = "Senza Nome"

                st.write(f"**Spazio di Lavoro**: {workspace_name} (ID: {sid})")

            with col2:
                if st.button("Usa", key=f"use_{sid}", use_container_width=True):
                    st.session_state.session_id = sid
                    st.success(f"Ora stai lavorando nello Spazio ID: {sid}")
                    load_forms_and_urls_data()
            with col3:
                if st.button("Elimina", key=f"delete_{sid}", use_container_width=True):
                    import shutil
                    full_path = get_user_session_dir(username, sid)
                    try:
                        shutil.rmtree(full_path)
                        st.success(f"Spazio di Lavoro {sid} eliminato con successo.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Errore durante l'eliminazione: {e}")

def chatbot_page():
    session_id = st.session_state.session_id
    st.sidebar.header("Report")

    if st.sidebar.button("Download Report", use_container_width=True):
        download_report()

    if st.sidebar.button("Genera Report", use_container_width=True):
        if session_id:
            st.sidebar.success("Generazione report in corso...")
            generate_report(session_id)
        else:
            st.sidebar.error("Per favore, inserisci un Session ID valido.")
    st.sidebar.markdown("---")

    st.sidebar.markdown(
        """
        <style>
        /* Contenitore con altezza fissa e scroll verticale */
        .fixed-height-chats {
            max-height: 250px; /* o l'altezza che preferisci */
            overflow-y: auto;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Sidebar for file upload and chain configuration
    #st.sidebar.header("Upload Documents")
    if st.session_state.logged_in:
        st.sidebar.subheader("Le tue chat esistenti")

        # 1) Pulsante "Crea Chat" con simbolo "+"
        if st.sidebar.button("➕ Crea Chat", use_container_width=True, key="new_chat_button"):
            new_id = str(uuid.uuid4())[:8]
            st.session_state.chat_id = new_id
            st.session_state.chat_name = "Nuova Chat"  # Nome di default
            st.session_state.messages = copy.deepcopy(chatbot_config["messages"])
            st.session_state.chat_history = copy.deepcopy(chatbot_config["messages"])
            st.rerun()

        # 2) Elenco chat in contenitore scrollabile
        user_chats_ids = list_user_chats(st.session_state["username"])

        with st.sidebar:  # Tutto ciò che segue viene visualizzato nella sidebar
            st.markdown("<div class='fixed-height-chats'>", unsafe_allow_html=True)

            # Se non esistono nel session_state, le inizializziamo
            if "show_rename_chat" not in st.session_state:
                st.session_state.show_rename_chat = None  # ID chat da rinominare (oppure None)
            if "rename_input_value" not in st.session_state:
                st.session_state.rename_input_value = ""

            user_chats_ids = list_user_chats(st.session_state["username"])

            for c_id in user_chats_ids:
                tmp_data = load_chat_from_file(st.session_state["username"], c_id)
                chat_name = tmp_data["name"]

                if c_id == st.session_state.get("chat_id"):
                    # CHAT CORRENTE: due colonne
                    colA, colB = st.columns([0.88, 0.12])
                    with colA:
                        if st.button(chat_name, use_container_width=True, key=f"load_chat_{c_id}"):
                            st.session_state.chat_id = c_id
                            st.session_state.chat_name = tmp_data["name"]
                            st.session_state.messages = tmp_data["messages"]
                            st.session_state.chat_history = copy.deepcopy(tmp_data["messages"])
                            st.rerun()
                    with colB:
                        # Pulsante ingranaggio per la chat attiva
                        if st.button("⚙", key=f"gear_chat_{c_id}"):
                            st.session_state.show_rename_chat = c_id
                            st.session_state.rename_input_value = chat_name
                else:
                    # CHAT NON CORRENTE: pulsante su tutta la larghezza
                    if st.button(chat_name, use_container_width=True, key=f"load_chat_{c_id}"):
                        st.session_state.chat_id = c_id
                        st.session_state.chat_name = tmp_data["name"]
                        st.session_state.messages = tmp_data["messages"]
                        st.session_state.chat_history = copy.deepcopy(tmp_data["messages"])
                        st.rerun()

                if st.session_state.show_rename_chat == c_id:
                    # Piccola spaziatura
                    #st.markdown("<div style='margin:5px 0;'></div>", unsafe_allow_html=True)

                    # Campo input SENZA etichetta, con placeholder
                    new_name = st.text_input(
                        label="",
                        value=st.session_state.rename_input_value,
                        key=f"rename_input_{c_id}",
                        placeholder="Inserisci nuovo nome..."
                    )

                    # Creiamo due colonne per affiancare 'Salva' e 'Elimina'
                    col_save, col_delete = st.columns(2)

                    with col_save:
                        if st.button("Salva", key=f"save_chat_{c_id}", use_container_width=True):
                            updated_chat_data = {
                                "id": c_id,
                                "name": new_name.strip() if new_name.strip() else "Nuova Chat",
                                "messages": tmp_data["messages"]
                            }
                            save_chat_to_file(st.session_state["username"], c_id, updated_chat_data)

                            # Se la chat rinominata è quella attiva, aggiorniamo anche session_state
                            if st.session_state.chat_id == c_id:
                                st.session_state.chat_name = updated_chat_data["name"]

                            # Nascondiamo form e ricarichiamo
                            st.session_state.show_rename_chat = None
                            st.success("Nome chat aggiornato!")
                            st.rerun()

                    with col_delete:
                        if st.button("Elimina", key=f"delete_chat_{c_id}", use_container_width=True):
                            # Eliminiamo fisicamente il file JSON corrispondente
                            chats_dir = get_user_chats_dir(st.session_state["username"])
                            chat_file = os.path.join(chats_dir, f"{c_id}.json")

                            try:
                                if os.path.exists(chat_file):
                                    os.remove(chat_file)
                                    # Se la chat eliminata è quella attiva, azzeriamo lo stato
                                    if st.session_state.chat_id == c_id:
                                        st.session_state.chat_id = None
                                        st.session_state.chat_name = None
                                        st.session_state.messages = []
                                        st.session_state.chat_history = []

                                    st.success("Chat eliminata con successo!")
                                else:
                                    st.warning("Il file della chat non esiste già.")
                            except Exception as e:
                                st.error(f"Errore durante l'eliminazione: {e}")

                            # Chiudiamo la form e ricarichiamo la sidebar
                            st.session_state.show_rename_chat = None
                            st.rerun()

                    # Piccola spaziatura
                    st.markdown("<div style='margin:5px 0;'></div>", unsafe_allow_html=True)

            st.markdown("</div>", unsafe_allow_html=True)

    # Sidebar: Add "Genera Report" button at the top
    #st.sidebar.header("Actions")
    #session_id = st.sidebar.text_input("Session ID", value="", placeholder="Inserisci Session ID")
    session_id = st.session_state.session_id
    # Sidebar for file upload and chain configuration

    st.sidebar.markdown("---")


    # Main chat interface
    previus_type = None
    for message in st.session_state.messages:

        if message["role"] == "user" and previus_type != "user":
            with st.chat_message(message["role"], avatar=st.session_state.user_avatar_url):
                st.markdown(message["content"])
        elif message["role"] != "user":
            with st.chat_message(message["role"], avatar=st.session_state.ai_avatar_url):
                st.markdown(message["content"])
        previus_type = message["role"]
    if prompt := st.chat_input("Scrivi qualcosa"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state.chat_history.append({"role": "user", "content": prompt})

        stop = False
        # Use auto_generated=True to avoid duplicating messages in chat history
        while not stop:
            try:
                generate_response(prompt, session_id, auto_generated=False)
                stop = True
            except Exception as e:
                print(f"Error: {e}")

    st.sidebar.markdown(
        """
        <div style="text-align: center; margin-top: 20px; font-size: 12px; color: black;">
            &copy; 2024 Bluen s.r.l. Tutti i diritti riservati.
        </div>
        """,
        unsafe_allow_html=True
    )


def format_questions_and_answers(data):
    """
    Elabora una lista di dizionari contenenti informazioni su domande e risposte.
    Per ogni elemento, restituisce una stringa formattata con:
      - La domanda (chiave 'domanda')
      - La risposta (chiave 'ANSWER')

    Se un elemento non contiene il campo 'ANSWER', viene usata una stringa predefinita.

    Args:
        data (list): Lista di dizionari con almeno le chiavi 'domanda' e 'ANSWER'.

    Returns:
        str: Testo formattato contenente tutte le domande e le relative risposte.
    """

    if not data:
        return ""

    lines = []
    for item in data:
        domanda = item.get('domanda', 'Domanda non disponibile')
        risposta = item.get('ANSWER', 'Nessuna risposta fornita')
        # Aggiungiamo una separazione fra le sezioni per chiarezza
        lines.append(f"Domanda: {domanda}")
        lines.append(f"Risposta: {risposta}")
        lines.append("-" * 80)  # separatore visivo
    return "\n".join(lines)


def generate_response(prompt, session_id, auto_generated=False):
    """Handles generating a response based on a prompt."""
    chain_id = f"{session_id}-workflow_generation_chain"

    forms_data = st.session_state.get("full_questionnaire_responses")
    print(format_questions_and_answers(forms_data))
    input_prefix = f"""
    ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    *NOTE IMPORTANTI:* 
    - Preleva informazioni utili dai vector stores, dai file in locale e dai seguenti URLs: {json.dumps(st.session_state.url_forms, indent=2) if st.session_state.get("url_forms") else "[urls assenti]"}
    - Inoltre dovrai usufruire delle informazioni ottrnute dai forms seguenti, ossia dai dati compilati in relazione alla azienda che si sta analizzando:
    
        '''
        {format_questions_and_answers(forms_data)}.
        '''
        
    - Preleva informazioni utili dai vector stores, dai file in locale e dai seguenti URL: {json.dumps(st.session_state.url_forms, indent=2) if st.session_state.get("url_forms") else "[urls assenti]"}
    - Inoltre dovrai usufruire delle informazioni ottrnute dai forms seguenti, ossia dai dati compilati in relazione alla azienda che si sta analizzando:
    ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    """
    # Show the user message in the chat only if it's not auto-generated
    if not auto_generated:
        with st.chat_message("user", avatar=st.session_state.user_avatar_url):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

    if auto_generated:
        st.session_state.chat_history.append({"role": "user", "content": prompt})

    with st.chat_message("assistant", avatar=st.session_state.ai_avatar_url):
        message_placeholder = st.empty()
        s = requests.Session()
        full_response = ""
        url = f"{api_address}/chains/stream_events_chain"
        payload = {
            "chain_id": chain_id,
            "query": {
                "input": f"{input_prefix}\n\n{prompt}",
                "chat_history": st.session_state.chat_history[-4:] if len(st.session_state.chat_history) > 4
                else st.session_state.chat_history
            },
            "inference_kwargs": {},
        }

        non_decoded_chunk = b''
        with s.post(url, json=payload, headers=None, stream=True) as resp:
            for chunk in resp.iter_content():
                if chunk:
                    non_decoded_chunk += chunk
                    if is_complete_utf8(non_decoded_chunk):
                        decoded_chunk = non_decoded_chunk.decode("utf-8")
                        full_response += decoded_chunk
                        message_placeholder.markdown(full_response + "▌", unsafe_allow_html=True)
                        non_decoded_chunk = b''  # Clear buffer

        message_placeholder.markdown(full_response)

    # Add assistant's response to chat history if not auto-generated
    st.session_state.messages.append({"role": "assistant", "content": full_response})
    st.session_state.chat_history.append({"role": "assistant", "content": full_response})

    if "username" in st.session_state and "chat_id" in st.session_state and "chat_name" in st.session_state:
        chat_data = {
            "id": st.session_state.chat_id,
            "name": st.session_state.chat_name,
            "messages": st.session_state.messages
        }
        save_chat_to_file(st.session_state["username"], st.session_state.chat_id, chat_data)

    aggiorna_html_address(full_response)
def generate_report(session_id):
    """Generates a report with predefined messages."""
    report_messages = [
    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Evita termini come 'benvenuti al' o 'nel nostro dna' etc..  Sviluppa contenuto sezione 'introduzione' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_intro}",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'introduzione' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave 'Introduzione' e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 1.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_1}",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli! In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inoltre cosa importantissima, assicurati di generare solo i grafici e le tabelle coerenti con l attuale sezione, ed evita la rindondanza di grafici identici ripetutti più volte! Inoltre se la sezione non richeide grafici allora non generarli! Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 1.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_2}",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 1.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_3}",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 1.4 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_4}",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.4 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 1.5 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_1_5}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 1.5 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto 'Descrizione Cap. 2' (nel prossimo messaggio ti chieiderò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Descrizione Cap. 2' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 2.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede .{linee_guida_2_1}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 2.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_2}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 2.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_3}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 2.4 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_4}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.4 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 2.5 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_5}.",
    # "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.5 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 2.6 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_2_6}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 2.6 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto 'Descrizione Cap. 3' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Descrizione Cap. 3' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 3.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3_1}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 3.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 3.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3_2}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 3.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 3.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_3_3}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 3.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto 'Descrizione Cap. 4' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_4}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Descrizione Cap. 4' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 4.1 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede.{linee_guida_4_1}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 4.1 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 4.2 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_4_2}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 4.2 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 4.3 (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_4_3}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 4.3 del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",


    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ...  contenuto sezione 'Conclusione' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_conclusione}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Conclusione' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 'Nota metodologica' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_nota_motedologica}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Nota metodologica' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 'Indice ESRS' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_indice_esrs}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Indice ESRS' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",

    f"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... Sviluppa contenuto sezione 'Informazioni di contatto' (nel prossimo messaggio ti chiederò di inserire il contenuto generato nel template html). Genera contenuton in formato html, duqnue inserisci contenuto dettagliato ed eventualmente inserisci tabelle, grafici e diagrammi quando la situazione lo richiede. {linee_guida_info_contatto}.",
    #"Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'  ... In base al contenuto appena generato per la sezione crea e salva effettivamente i grafici richiesti, duqnue riscrivi il contenuto html inserendo i link che puntano ai grafici salvati. I grafici salvati andranno visualizzati nell html della sezione centrati orizzontalmente.",
    "Il contenuto che svilupperai dovrà avere un stilo infromale e professionale, inoltre riporta solo dati e informazioni attendibili estratte dalle risorse fornite e inerenti l'azienda analizzata.Evita termini come 'benvenuti al' o 'nel nostro dna' etc.. Evita frasi del tipo 'Non sono riuscito a ottenere informazioni dettagliate dai vector stores'   ... Inserisci contenuto generato nel messaggio precedente all'interno della sezione 'Informazioni di contatto' del template. USA assolutamente STRUMENTO replace_placeholder_in_html fornendo in input chiave e contenuti adatti in formato html",
]


    for message in report_messages:
        stop = False
        # Use auto_generated=True to avoid duplicating messages in chat history
        while not stop:
            try:
                generate_response(message, session_id, auto_generated=True)
                stop = True

            except Exception as e:
                print(f"Error: {e}")

    # AL TERMINE: salvataggio automatico del docx
    username = st.session_state.get("username")
    if username and st.session_state.current_html_address:
        html_content = scarica_html(st.session_state.current_html_address)
        if html_content:
            docx_file_path = "final_report.docx"  # Nome di base, volendo puoi usare un timestamp
            docx_file = converti_html_in_docx(html_content, docx_file_path)

            # Ottieni la directory dell'utente per i report
            user_reports_dir = get_user_reports_dir(username)

            # Scegli un nome univoco (puoi usare un timestamp, un uuid, ecc.)
            unique_filename = f"report_{uuid.uuid4().hex[:8]}.docx"
            full_save_path = os.path.join(user_reports_dir, unique_filename)

            # Salva effettivamente il file nella cartella utente
            if docx_file:
                # docx_file qui è "report.docx" già scritto su disco,
                # se la tua funzione `converti_html_in_docx` crea un file fisico.
                # In caso tu stia restituendo `bytes`, cambia di conseguenza.
                if os.path.exists(docx_file):
                    os.rename(docx_file, full_save_path)
                else:
                    # Se converti_html_in_docx restituisce bytes, gestisci così:
                    # with open(full_save_path, "wb") as f:
                    #     f.write(docx_file)
                    pass

                st.success(f"Report generato e salvato come: {full_save_path}")
            else:
                st.error("Impossibile convertire il contenuto HTML in DOCX.")
        else:
            st.error("Impossibile scaricare l'HTML dall'indirizzo corrente.")
    else:
        st.error("Report non generato o utente non loggato.")

########################################################################################################################

def main():
    """
    Gestisce la logica principale dell'app Streamlit.
    """
    # Sidebar per la navigazione tra Chatbot e Questionario

    #col1, col2 = st.sidebar.columns(2)

    st.sidebar.markdown(
        """
        <div style="text-align: center;">
            <img src="https://static.wixstatic.com/media/63b1fb_6c4848f63b21475db8775af35ee50e55~mv2.png" alt="Logo" style="width: 300px;">
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.markdown("---")

    # Mostra il nome dello Spazio di Lavoro corrente (se esiste)
    if st.session_state.get("username") and st.session_state.get("session_id"):
        session_path = get_user_session_dir(st.session_state["username"], st.session_state["session_id"])
        name_file = os.path.join(session_path, "workspace_name.txt")
        if os.path.exists(name_file):
            with open(name_file, "r", encoding="utf-8") as f:
                workspace_name = f.read().strip()
            st.sidebar.markdown(f"**Spazio di Lavoro Attuale:** {workspace_name}")
        else:
            st.sidebar.markdown("*Nessuno Spazio di Lavoro selezionato.*")
    else:
        st.sidebar.markdown("*Nessuno Spazio di Lavoro selezionato.*")

    st.sidebar.header("Pages")

    with st.sidebar:
        if st.button("Bluen AI", key="chatbot_button", use_container_width=True):
            st.session_state["current_page"] = "chatbot"

        # AGGIUNGI QUESTO TASTO
        if st.sidebar.button("Spazi di Lavoro", key="workspace_button", use_container_width=True):
            st.session_state["current_page"] = "workspace"

        #with col2:
        if st.button("ESG Assessment Form", key="questionnaire_button", use_container_width=True):
            st.session_state["current_page"] = "questionnaire"
        #-
        # AGGIUNTA: bottone per la pagina Scope 2
        if st.sidebar.button("GHG Form (Scope 1)", key="scope1_button", use_container_width=True):
            st.session_state["current_page"] = "scope1"

        # AGGIUNTA: bottone per la pagina Scope 2
        if st.sidebar.button("GHG Form (Scope 2)", key="scope2_button", use_container_width=True):
            st.session_state["current_page"] = "scope2"

        # AGGIUNTA: bottone per la pagina Scope 2
        if st.sidebar.button("GHG Form (Scope 3)", key="scope3_button", use_container_width=True):
            st.session_state["current_page"] = "scope3"

        if st.button("Documents", key="documents_button", use_container_width=True):
            st.session_state["current_page"] = "documents"

        if st.sidebar.button("Reports", key="reports_button", use_container_width=True):
            st.session_state["current_page"] = "reports"

        # Se l'utente loggato è admin, mostra il bottone per la Dashboard
        if st.session_state.get("username") == "admin":
            if st.sidebar.button("Dashboard", key="dashboard_button", use_container_width=True):
                st.session_state["current_page"] = "dashboard"

    st.sidebar.markdown("---")

    if st.session_state.get("current_page", "chatbot") == "chatbot":
        chatbot_page()  # Funzione che gestisce la logica del chatbot
    elif st.session_state.get("current_page") == "questionnaire":
        questionnaire_page()
    #-
    elif st.session_state.get("current_page") == "scope1":
        render_scope1_form(session_id=st.session_state.session_id)
    elif st.session_state.get("current_page") == "scope2":
        render_scope2_form(session_id=st.session_state.session_id)
    elif st.session_state.get("current_page") == "scope3":
        render_scope3_form(session_id=st.session_state.session_id)
    elif st.session_state.get("current_page") == "dashboard":  # <<-- AGGIUNGI QUESTA CONDIZIONE
        dashboard_page()
    elif st.session_state.get("current_page") == "documents":
        documents_page()
    elif st.session_state.get("current_page") == "reports":
        reports_page()
    elif st.session_state["current_page"] == "workspace":
        workspace_management_page()

# Se l'utente non è loggato, mostra la login_page
if not st.session_state.logged_in:
    login_page()
else:
    # Se loggato, richiama main()
    main()


