from docx import Document
from html2docx import html2docx
import re

html_content = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Bilancio di Sostenibilità 2024</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 20px;
            color: #333;
        }
        h1, h2, h3 {
            color: #0056b3;
        }
        p {
            line-height: 1.5;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }
        table, th, td {
            border: 1px solid #ddd;
        }
        th, td {
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f4f4f4;
        }
        .highlight {
            background-color: #eaf4ff;
        }
    </style>
</head>
<body>
    <h1>Bilancio di Sostenibilità</h1>
    <h2>Introduzione</h2>
    <p>Questo documento rappresenta un esempio di bilancio di sostenibilità. Include contenuti dimostrativi per testare il sistema di conversione HTML in DOCX.</p>

    <h2>Tabella di Esempio</h2>
    <table>
        <thead>
            <tr>
                <th>Voce</th>
                <th>Descrizione</th>
                <th>Valore</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td class="highlight">Emissioni CO2</td>
                <td>Totale emissioni per l'anno 2024</td>
                <td>1000 t</td>
            </tr>
            <tr>
                <td>Consumo Energetico</td>
                <td>Energia utilizzata (kWh)</td>
                <td>500.000 kWh</td>
            </tr>
        </tbody>
    </table>

    <h2>Conclusione</h2>
    <p>Questo esempio conclude il test. Grazie per aver utilizzato il nostro sistema!</p>
</body>
</html>
"""

def converti_html_in_docx_con_carta_intestata(html_content, template_path, output_path):
    """
    Converte l'HTML in contenuto DOCX e lo aggiunge a un documento già esistente con carta intestata.

    Args:
        html_content (str): Contenuto HTML da convertire.
        template_path (str): Percorso del documento DOCX esistente con carta intestata.
        output_path (str): Percorso in cui salvare il documento aggiornato.
    """
    try:
        # Rimuovi il blocco <head> dal contenuto HTML
        pattern = r"<head>.*?</head>"
        html_content = re.sub(pattern, "", html_content, flags=re.DOTALL)

        # Converti l'HTML in un contenuto DOCX (come bytes)
        docx_content = html2docx(title="Document from HTML", content=html_content).getvalue()

        # Apri il documento DOCX esistente con carta intestata
        template_doc = Document(template_path)

        # Crea un documento temporaneo per il contenuto HTML convertito
        temp_doc = Document()
        with open("temp_doc.docx", "wb") as temp_file:
            temp_file.write(docx_content)
        temp_doc = Document("temp_doc.docx")

        # Aggiungi il contenuto HTML convertito al documento esistente
        for element in temp_doc.element.body:
            template_doc.element.body.append(element)

        # Salva il documento aggiornato con il contenuto HTML
        template_doc.save(output_path)
        print(f"Documento salvato con successo in: {output_path}")
    except Exception as e:
        print(f"Errore durante la conversione o l'inserimento: {e}")


converti_html_in_docx_con_carta_intestata(html_content=html_content, template_path="Carta_intestata_Bluen.docx", output_path="output.docx")

